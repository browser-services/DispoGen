import random
from datetime import datetime
import sqlite3
import os
from telegram import InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, CommandHandler, CallbackQueryHandler, MessageHandler, filters
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

# --- CONFIG ---
TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
DB_PATH = "users_reports.db"
TEMPLATE_DOCX = "template.docx"

# --- Helper Functions ---
def fetch_users():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT id, full_name FROM Users")
    users = cur.fetchall()
    conn.close()
    return users

def fetch_reports(user_id):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT text FROM Reports WHERE user_id=?", (user_id,))
    reports = [row[0] for row in cur.fetchall()]
    conn.close()
    return reports

def generate_docx(user_fullname, reports):
    from datetime import datetime, timedelta, timezone

    # --- Philippine timezone ---
    PH_TZ = timezone(timedelta(hours=8))

    # Yesterday in PH time
    today_ph = datetime.now(PH_TZ)
    yesterday_ph = today_ph - timedelta(days=1)

    day = yesterday_ph.strftime("%d")
    month = yesterday_ph.strftime("%B")
    year = yesterday_ph.strftime("%Y")

    doc = Document(TEMPLATE_DOCX)
    table = doc.tables[0]

    # Name
    cell = table.cell(0, 0)
    cell.text = f"Name: {user_fullname}"
    run = cell.paragraphs[0].runs[0]
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(13)

    # DISPOSITION Header
    cell = table.cell(1, 1)
    cell.text = ""  # clear existing paragraphs

    # DISPOSITION text
    para1 = cell.paragraphs[0]
    para1.text = "DISPOSITION"
    run1 = para1.runs[0]
    run1.font.name = "Arial"
    run1.font.bold = True
    run1.font.size = Pt(16)
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Covered Period
    para2 = cell.add_paragraph(f"(Covered Period:{day}0800 – 2000 {month} {year})")
    run2 = para2.runs[0]
    run2.font.name = "Arial"
    run2.font.bold = True
    run2.font.italic = True
    run2.font.size = Pt(13)
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Reports
    max_reports = min(3, len(reports))
    if max_reports > 0:
        selected_reports = random.sample(reports, k=max_reports)
        allowed_rows = list(range(3, 15))
        used_rows = set()
        for report in selected_reports:
            span = random.choice([2, 3, 4])
            possible_starts = [
                r for r in allowed_rows
                if r + span - 1 <= 14 and not any(rr in used_rows for rr in range(r, r + span))
            ]
            if not possible_starts:
                break
            start_row = random.choice(possible_starts)
            for rr in range(start_row, start_row + span):
                used_rows.add(rr)
            start_cell = table.cell(start_row, 2)
            end_cell = table.cell(start_row + span - 1, 2)
            merged_cell = start_cell.merge(end_cell)
            merged_cell.text = ""
            para = merged_cell.paragraphs[0]
            run = para.add_run(report)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            template_run = table.cell(1, 3).paragraphs[0].runs[0]
            run.font.name = template_run.font.name
            run.font.size = template_run.font.size
            para.paragraph_format.space_before = 0
            para.paragraph_format.space_after = 0

    # Filepath using yesterday's date
    filepath = f"DISPOSITION_{day}_{month}_{year}_{user_fullname.split()[0]}.docx"
    doc.save(filepath)
    return filepath

# --- Safe send/edit ---
async def safe_send(bot, chat_id, text, reply_markup=None):
    try:
        await bot.send_message(chat_id=chat_id, text=text, reply_markup=reply_markup)
    except:
        pass

async def safe_edit(bot, chat_id, message_id, text, reply_markup=None):
    try:
        await bot.edit_message_text(chat_id=chat_id, message_id=message_id, text=text, reply_markup=reply_markup)
    except:
        pass

def get_chat_id(update_or_query):
    if hasattr(update_or_query, "effective_chat") and update_or_query.effective_chat:
        return update_or_query.effective_chat.id
    elif hasattr(update_or_query, "callback_query") and update_or_query.callback_query:
        return update_or_query.callback_query.message.chat.id
    else:
        return None

# --- Menu ---
async def show_menu(update_or_query, context):
    keyboard = [
        [InlineKeyboardButton("Generate Report", callback_data="menu_generate")],
        [InlineKeyboardButton("View Activities", callback_data="menu_view")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    chat_id = get_chat_id(update_or_query)
    if chat_id is None:
        return
    if hasattr(update_or_query, "message") and update_or_query.message:
        await safe_send(context.bot, chat_id, "Select an option:", reply_markup=reply_markup)
    elif hasattr(update_or_query, "callback_query") and update_or_query.callback_query:
        msg_id = update_or_query.callback_query.message.message_id
        await safe_edit(context.bot, chat_id, msg_id, "Select an option:", reply_markup=reply_markup)

# --- Start ---
async def start(update, context):
    await show_menu(update, context)

# --- Menu Callback ---
async def menu_callback(update, context):
    query = update.callback_query
    await query.answer()
    data = query.data

    if data == "menu_generate":
        await generate_report_start(update, context)
    elif data == "menu_view":
        users = fetch_users()
        keyboard = [[InlineKeyboardButton(u[1], callback_data=f"view_{u[0]}")] for u in users]
        keyboard.append([InlineKeyboardButton("Back", callback_data="menu_back")])
        reply_markup = InlineKeyboardMarkup(keyboard)
        await safe_edit(context.bot, query.message.chat.id, query.message.message_id,
                        "Select a user to view activities:", reply_markup=reply_markup)
    elif data.startswith("view_"):
        await view_user_options(update, context)
    elif data.startswith("add_") or data.startswith("remove_"):
        await activity_user_selected(update, context)
    elif data == "menu_back":
        await show_menu(update, context)


# --- View Activities ---
async def view_user_options(update, context):
    query = update.callback_query
    await query.answer()
    user_id = int(query.data.split("_")[1])
    context.user_data["user_id"] = user_id
    reports = fetch_reports(user_id)
    user_fullname = [u[1] for u in fetch_users() if u[0] == user_id][0]

    if not reports:
        msg = f"No reports for {user_fullname}."
    else:
        msg = f"Activities for {user_fullname}:\n"
        for idx, r in enumerate(reports, 1):
            msg += f"{idx}. {r}\n"

    keyboard = [
        [InlineKeyboardButton("Add Activity", callback_data=f"add_{user_id}")],
        [InlineKeyboardButton("Remove Activity", callback_data=f"remove_{user_id}")],
        [InlineKeyboardButton("Back", callback_data="menu_back")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await safe_edit(context.bot, query.message.chat.id, query.message.message_id, msg, reply_markup=reply_markup)

# --- Generate Report ---
async def generate_report_start(update, context):
    users = fetch_users()
    keyboard = [[InlineKeyboardButton(u[1], callback_data=f"report_{u[0]}")] for u in users]
    reply_markup = InlineKeyboardMarkup(keyboard)
    chat_id = update.effective_chat.id
    await safe_send(context.bot, chat_id, "Select a user:", reply_markup=reply_markup)

async def report_user_selected(update, context):
    query = update.callback_query
    await query.answer()
    user_id = int(query.data.split("_")[1])
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT full_name FROM Users WHERE id=?", (user_id,))
    user_fullname = cur.fetchone()[0]
    conn.close()
    reports = fetch_reports(user_id)
    if not reports:
        await safe_edit(context.bot, query.message.chat.id, query.message.message_id,
                        f"No reports found for {user_fullname}.")
        await show_menu(query, context)
        return
    filepath = generate_docx(user_fullname, reports)
    with open(filepath, "rb") as f:
        await context.bot.send_document(chat_id=query.message.chat.id, document=f, filename=filepath)
    os.remove(filepath)
    await safe_edit(context.bot, query.message.chat.id, query.message.message_id,
                    f"Report generated for {user_fullname}.")
    await show_menu(query, context)

# --- Add / Remove Activity ---
async def activity_user_selected(update, context):
    query = update.callback_query
    await query.answer()
    data = query.data
    user_id = int(data.split("_")[1])
    context.user_data["user_id"] = user_id
    context.user_data["pending_action"] = "add" if data.startswith("add_") else "remove"
    reports = fetch_reports(user_id)

    if context.user_data["pending_action"] == "remove":
        if not reports:
            await safe_edit(context.bot, query.message.chat.id, query.message.message_id,
                            "No reports to remove for this user.")
            await show_menu(query, context)
            return
        msg = "Send the numbers of reports to remove, separated by commas:\n"
        for idx, r in enumerate(reports, 1):
            msg += f"{idx}. {r}\n"
        await safe_edit(context.bot, query.message.chat.id, query.message.message_id, msg)
    else:
        await safe_edit(context.bot, query.message.chat.id, query.message.message_id,
                        "Send the new report(s) to add. Separate multiple reports by new lines.")

# --- Handle Messages ---
async def handle_activity_message(update, context):
    action = context.user_data.get("pending_action")
    user_id = context.user_data.get("user_id")
    if not action or not user_id:
        return
    text = update.message.text.strip()
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    if action == "add":
        reports = [r.strip() for r in text.split("\n") if r.strip()]
        for r in reports:
            cur.execute("INSERT INTO Reports (user_id, text) VALUES (?, ?)", (user_id, r))
        conn.commit()
        await update.message.reply_text(f"{len(reports)} report(s) added successfully.")
    elif action == "remove":
        indices = [int(i.strip()) for i in text.split(",") if i.strip().isdigit()]
        reports = fetch_reports(user_id)
        removed_count = 0
        for i in indices:
            if 1 <= i <= len(reports):
                cur.execute("DELETE FROM Reports WHERE user_id=? AND text=?", (user_id, reports[i-1]))
                removed_count += 1
        conn.commit()
        await update.message.reply_text(f"{removed_count} report(s) removed successfully.")
    conn.close()
    context.user_data.pop("pending_action", None)
    context.user_data.pop("user_id", None)
    await show_menu(update, context)

# --- Main ---
app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()

app.add_handler(CommandHandler("start", start))
app.add_handler(CallbackQueryHandler(menu_callback, pattern="^menu_"))
app.add_handler(CallbackQueryHandler(report_user_selected, pattern="^report_"))
app.add_handler(CallbackQueryHandler(activity_user_selected, pattern="^(add_|remove_)"))
app.add_handler(CallbackQueryHandler(view_user_options, pattern="^view_"))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_activity_message))

app.run_polling()
